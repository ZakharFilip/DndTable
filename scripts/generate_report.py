# -*- coding: utf-8 -*-
"""Generate ОтчётЗаписка_полный.md and ОтчётЗаписка.docx (GOST formatting via python-docx)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_OUT = ROOT / "ОтчётЗаписка_полный.md"
DOCX_OUT = ROOT / "ОтчётЗаписка.docx"

BIBLIOGRAPHY = """
1. Васильев, П. П. Информационные системы и технологии : учебник / П. П. Васильев. — М. : Юрайт, 2022. — 384 с.

2. Гагарина, Л. Г. Разработка и эксплуатация автоматизированных информационных систем : учеб. пособие / Л. Г. Гагарина, В. Д. Кокорева, Б. Д. Виснадул. — М. : ИД «Форум» : ИНФРА-М, 2019. — 384 с.

ГОСТ 7.32-2017. Отчёт о научно-исследовательской работе. Структура и правила оформления. — М. : Стандартинформ, 2018.

ГОСТ 34.602-89. Техническое задание. Требования к содержанию и оформлению.

Дейт, К. Дж. Введение в системы баз данных / К. Дж. Дейт ; пер. с англ. — 8-е изд. — М. : Вильямс, 2020. — 1328 с.

Коннолли, Т. Базы данных. Проектирование, реализация и сопровождение. Теория и практика / Т. Коннолли, К. Бегг ; пер. с англ. — 3-е изд. — М. : Вильямс, 2016. — 1440 с.

Макконнелл, С. Совершенный код / С. Макконнелл ; пер. с англ. — М. : Русская редакция, 2019. — 896 с.

Мартин, Р. Чистая архитектура. Искусство разработки программного обеспечения / Р. Мартин ; пер. с англ. — СПб. : Питер, 2018. — 352 с.

Орлов, С. А. Технологии разработки программного обеспечения : учебник / С. А. Орлов. — 4-е изд. — СПб. : Питер, 2019. — 464 с.

Роганов, Е. А. UML 2.5. Руководство пользователя / Е. А. Роганов. — СПб. : БХВ-Петербург, 2020. — 304 с.

Соммервилл, И. Инженерия программного обеспечения / И. Соммервилл ; пер. с англ. — 10-е изд. — М. : Вильямс, 2016. — 624 с.

Таненбаум, Э. Компьютерные сети / Э. Таненбаум, Д. Уэзеролл ; пер. с англ. — 5-е изд. — СПб. : Питер, 2018. — 960 с.

Фаулер, М. Архитектура корпоративных программных приложений / М. Фаулер ; пер. с англ. — М. : Вильямс, 2019. — 544 с.

Хомоненко, А. Д. Базы данных : учебник / А. Д. Хомоненко, М. Г. Дубейковский, В. М. Колесников. — 10-е изд. — СПб. : Корона-Век, 2019. — 736 с.

Шилдт, Г. Node.js. Подробное руководство / Г. Шилдт ; пер. с англ. — 2-е изд. — М. : Вильямс, 2020. — 640 с.

Banks, A. Learning React : Modern Patterns for Developing React Apps / A. Banks, E. Porcello. — 2nd ed. — O'Reilly Media, 2020. — 310 p.

Booch, G. Object-Oriented Analysis and Design with Applications / G. Booch et al. — 3rd ed. — Addison-Wesley, 2007. — 720 p.

MongoDB Documentation [Электронный ресурс]. — URL: https://www.mongodb.com/docs/ (дата обращения: 16.05.2026).

Node.js Documentation [Электронный ресурс]. — URL: https://nodejs.org/docs/ (дата обращения: 16.05.2026).

React Documentation [Электронный ресурс]. — URL: https://react.dev/ (дата обращения: 16.05.2026).

Roll20 — виртуальный стол для настольных ролевых игр [Электронный ресурс]. — URL: https://roll20.net/ (дата обращения: 16.05.2026).

Foundry Virtual Tabletop [Электронный ресурс]. — URL: https://foundryvtt.com/ (дата обращения: 16.05.2026).

Owlbear Rodeo [Электронный ресурс]. — URL: https://www.owlbear.rodeo/ (дата обращения: 16.05.2026).

Socket.IO Documentation [Электронный ресурс]. — URL: https://socket.io/docs/v4/ (дата обращения: 16.05.2026).

TypeScript Handbook [Электронный ресурс]. — URL: https://www.typescriptlang.org/docs/ (дата обращения: 16.05.2026).

Zod Documentation [Электронный ресурс]. — URL: https://zod.dev/ (дата обращения: 16.05.2026).

Виртуальные столы для D&D: обзор рынка VTT [Электронный ресурс] // EN World. — URL: https://www.enworld.org/ (дата обращения: 16.05.2026).

ГОСТ Р ИСО/МЭК 25010-2015. Системная и программная инженерия. Требования и оценка качества систем и программного обеспечения.
""".strip()


def strip_citations(text: str) -> str:
    """Remove bibliographic references like [8], [12, с. 44], [19; 20]."""
    text = re.sub(r"\s*\[\d+(?:,\s*с\.\s*\d+)?(?:;\s*\d+)*\]", "", text)
    return text


def build_module_narratives() -> str:
    """Supplementary implementation detail for code-backed modules only."""
    return """
#### 3.1.7 Детализация каталога backend/src/modules/auth

Модуль размещён в `backend/src/modules/auth`. Точка входа — `auth.router.ts`, бизнес-логика — `auth.service.ts`. Регистрация принимает email, username, password. Валидация express-validator проверяет формат почты, длину никнейма и минимальную длину пароля. Сервис проверяет уникальность email и username, применяет `_isPasswordStrong`, хеширует пароль bcrypt. При успехе в сессию записывается userId. Middleware `requireAuth` отклоняет запросы без авторизации кодом 401 JSON.

#### 3.1.8 Детализация каталога backend/src/modules/gamesessions

`GameSessionsService` инкапсулирует создание сессии, выборку списков, загрузку полного состояния и применение патчей. Модели `GameSession`, `TableObject`, `SessionState` описаны в Mongoose-схемах. Уникальный индекс по паре `(gameSessionId, key)` гарантирует идемпотентность создания объектов с клиентским ключом.

#### 3.1.9 Детализация каталога frontend/src/tabletop

Подсистема включает `geometry.ts`, `CanvasRenderer`, `TableController`, `TableSync`, `HistoryManager`. Vitest-тесты выполняются без поднятия DOM в headless-режиме. Hook `useTableSync` управляет жизненным циклом соединения и flush при `beforeunload`.
"""


def build_markdown() -> str:
    parts: list[str] = []

    parts.append("""# ПОЯСНИТЕЛЬНАЯ ЗАПИСКА

к дипломному проекту на тему: «Разработка веб-системы виртуального стола DnDTable для проведения настольных ролевых игр»

---

## СОДЕРЖАНИЕ

(В документе Word оформляется автоматическим оглавлением по заголовкам 1–3 уровня.)

Введение

1 Аналитическая часть

1.1 Описание предметной области

1.2 Анализ существующих аналогов

1.3 Требования к разрабатываемой информационной системе

1.3.1 Функциональные требования

1.3.2 Нефункциональные требования

1.4 Анализ методов решения задачи

1.5 Обоснование выбора методики, технологии и инструментальных средств проектирования и разработки

2 Проектная часть

2.1 Разработка модели базы данных

2.2 Проектирование программного обеспечения информационной системы

2.3 Проектирование интерфейсов

3 Реализация и тестирование программного обеспечения

3.1 Реализация основных функций программного обеспечения

3.2 Тестирование программного обеспечения

Заключение

Перечень принятых терминов и сокращений

Список использованных источников

Приложение Б Диаграммы UML (текстовое описание)

---

## ВВЕДЕНИЕ

Современный этап развития общества характеризуется высокой степенью информатизации социальных и досуговых практик. Средства информационных технологий проникают в сферы, где ещё недавно доминировали исключительно офлайн-форматы взаимодействия. К таким сферам относится и организация настольных ролевых игр (TTRPG): группы игроков всё чаще используют сетевые сервисы для совместного ведения партий, отображения карт, перемещения фигур и обмена игровыми материалами.

Виртуальный стол (Virtual Tabletop, VTT) представляет собой программный комплекс, имитирующий игровое поле, объекты на нём и правила взаимодействия участников. Качество VTT определяется не только скоростью отрисовки, но и согласованностью данных между клиентами, возможностью разграничения прав, удобством подготовки сцен и надёжностью хранения пользовательских материалов. Коммерческие и открытые решения на рынке демонстрируют различные компромиссы между простотой, функциональностью и стоимостью владения.

Актуальность разработки веб-системы DnDTable обусловлена следующими факторами. Во-первых, распространённые облачные платформы ограничивают гибкость настройки прав доступа и хранение данных на стороне пользователя. Во-вторых, самостоятельно разворачиваемые решения нередко требуют сложной инсталляции и экосистемы модулей, что повышает порог входа для ведущего игры. В-третьих, целесообразно иметь прозрачный стек TypeScript с единым контрактом данных между клиентом и сервером, что упрощает сопровождение и развитие системы.

Это послужило основанием для выбора темы выпускной квалификационной работы: разработка веб-ориентированной автоматизированной информационной системы **DnDTable** — виртуального двумерного стола для проведения настольных ролевых игр с поддержкой партий, сцен, компонентной модели объектов, синхронизации в реальном времени и разграничения доступа.

**Цель работы** — разработать информационную систему DnDTable, обеспечивающую совместную работу ведущего и игроков над игровыми сценами, централизованное хранение состояния, контроль прав и обмен изменениями в режиме реального времени.

Для достижения указанной цели поставлены следующие **задачи**:

1. Изучить предметную область настольных ролевых игр и требования к виртуальным столам.
2. Проанализировать существующие аналоги и обосновать необходимость разработки собственной системы.
3. Сформировать функциональные и нефункциональные требования к DnDTable.
4. Разработать информационную модель данных и архитектуру программного обеспечения.
5. Спроектировать пользовательские интерфейсы ключевых сценариев (авторизация, партия, редактор сцены).
6. Реализовать программные модули серверной и клиентской частей, общий пакет схем валидации.
7. Реализовать механизмы синхронизации состояния сцены и контроля доступа.
8. Провести тестирование системы и оценить соответствие требованиям.

**Объект исследования** — процессы организации и проведения совместных игровых сессий в цифровой среде.

**Предмет исследования** — методы и средства проектирования веб-ориентированной информационной системы виртуального стола.

**Методы исследования** — анализ предметной области и аналогов, объектно-ориентированное проектирование, прототипирование архитектуры «клиент — сервер», экспериментальная проверка (тестирование).

---

## 1 Аналитическая часть

### 1.1 Описание предметной области
""")

    # Continue building in append function - read from external file for size
    body_path = Path(__file__).with_name("report_body.md")
    exp_path = Path(__file__).with_name("report_expansion.md")
    api_path = Path(__file__).with_name("report_api_tests.md")
    academic_path = Path(__file__).with_name("report_academic.md")
    appendix_b_path = Path(__file__).with_name("report_appendix_b.md")

    chunks = [
        body_path.read_text(encoding="utf-8"),
        exp_path.read_text(encoding="utf-8"),
        academic_path.read_text(encoding="utf-8") if academic_path.exists() else "",
        api_path.read_text(encoding="utf-8"),
        build_module_narratives(),
    ]
    parts.append(strip_citations("\n".join(chunks)))

    parts.append(strip_citations(f"""

---

## ЗАКЛЮЧЕНИЕ

В ходе выполнения выпускной квалификационной работы разработана информационная система DnDTable — веб-ориентированный виртуальный стол для настольных ролевых игр. Поставленная цель достигнута: спроектирована и реализована архитектура, обеспечивающая совместную работу участников над игровым столом, хранение данных на сервере, синхронизацию изменений и проектирование разграничения прав.

В аналитической части описана предметная область, проведён сравнительный анализ аналогов Roll20, Foundry VTT и Owlbear Rodeo, сформулированы функциональные и нефункциональные требования. Обоснован выбор клиент-серверной архитектуры, MongoDB, Node.js/Express, React, Socket.IO и компонентной модели объектов. Зафиксированы принятые технологические решения: серверные сессии, Canvas 2D, протокол патчей с оптимистичной конкуренцией.

В проектной части разработана логическая модель базы данных, архитектура модулей монорепозитория, варианты использования и структура пользовательского интерфейса редактора. Текстовые описания диаграмм приведены в приложении Б.

В разделе реализации и тестирования описаны модули авторизации, игровых сессий, realtime-шлюза и клиентского редактора; проведены ручные и автоматизированные испытания. Критерии приёмки согласованы с результатами тестирования.

Перспективы развития связаны с эксплуатационным качеством: развёртывание кластера Node.js с Redis-адаптером Socket.IO для горизонтального масштабирования; подключение объектного хранилища S3/MinIO для ассетов; внедрение мониторинга (Prometheus/Grafana) и rate limiting на маршрутах авторизации; расширение набора интеграционных и нагрузочных тестов; периодический аудит безопасности веб-приложения.

---

## ПЕРЕЧЕНЬ ПРИНЯТЫХ ТЕРМИНОВ И СОКРАЩЕНИЙ

**ACL** (Access Control List) — список управления доступом; набор правил, определяющих разрешённые действия субъекта над ресурсом.

**API** (Application Programming Interface) — программный интерфейс приложения.

**ECS** (Entity–Component System) — подход к представлению объектов сцены как сущностей с набором компонентов.

**GM** (Game Master) — ведущий игры.

**REST** — архитектурный стиль взаимодействия компонентов через HTTP-методы.

**SPA** (Single Page Application) — одностраничное веб-приложение.

**TTRPG** (Tabletop Role-Playing Game) — настольная ролевая игра.

**VTT** (Virtual Tabletop) — виртуальный стол.

**WebSocket** — протокол полнодуплексной связи поверх TCP для обмена сообщениями в реальном времени.

---

## СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ

{BIBLIOGRAPHY}

---

"""))
    if appendix_b_path.exists():
        parts.append(strip_citations(appendix_b_path.read_text(encoding="utf-8")))

    return strip_citations("\n".join(parts))


# --- DOCX (GOST via python-docx) ---

STRUCTURAL_HEADS = frozenset(
    {
        "СОДЕРЖАНИЕ",
        "ВВЕДЕНИЕ",
        "ЗАКЛЮЧЕНИЕ",
        "ПЕРЕЧЕНЬ ПРИНЯТЫХ ТЕРМИНОВ И СОКРАЩЕНИЙ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    }
)
PAGE_BREAK_PREFIXES = (
    "ВВЕДЕНИЕ",
    "ЗАКЛЮЧЕНИЕ",
    "1 АНАЛИТИЧЕСКАЯ",
    "2 ПРОЕКТНАЯ",
    "3 РЕАЛИЗАЦИЯ",
    "ПЕРЕЧЕНЬ",
    "СПИСОК",
    "ПРИЛОЖЕНИЕ",
)


def _clean_md_text(text: str) -> str:
    """Убрать разметку Markdown (*, **, `)."""
    s = text.strip()
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"(?<!\*)\*([^*\s][^*]*?)\*(?!\*)", r"\1", s)
    s = s.replace("**", "").replace("*", "")
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s.strip()


def _set_run_font(run, name: str = "Times New Roman", size_pt: float = 14, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.insert(0, rFonts)


def _setup_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    normal = doc.styles["Normal"]
    nf = normal.font
    nf.name = "Times New Roman"
    nf.size = Pt(14)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for lvl, size in (("Heading 1", 14), ("Heading 2", 14), ("Heading 3", 14)):
        if lvl not in doc.styles:
            continue
        st = doc.styles[lvl]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(6)
        st.paragraph_format.space_after = Pt(6)


class _DocState:
    __slots__ = ("has_content", "pending_page_break")

    def __init__(self) -> None:
        self.has_content = False
        self.pending_page_break = False


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _request_page_break(state: _DocState) -> None:
    if state.has_content:
        state.pending_page_break = True


def _flush_page_break(doc: Document, state: _DocState) -> None:
    if state.pending_page_break:
        _add_page_break(doc)
        state.pending_page_break = False


def _mark_content(state: _DocState) -> None:
    state.has_content = True


def _add_page_number_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)
    _set_run_font(run, size_pt=14)


def _add_toc(doc: Document, state: _DocState) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_placeholder = OxmlElement("w:t")
    fld_placeholder.text = "Оглавление (обновите поле)"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_placeholder)
    run._r.append(fld_end)
    _set_run_font(run, size_pt=14)
    hint = doc.add_paragraph(
        "ПКМ по оглавлению → «Обновить поле» → «Обновить целиком»."
    )
    hint.paragraph_format.first_line_indent = Cm(0)
    for r in hint.runs:
        _set_run_font(r, size_pt=12)
    _mark_content(state)


def _para_structural(doc: Document, text: str, state: _DocState) -> None:
    _flush_page_break(doc, state)
    p = doc.add_paragraph(_clean_md_text(text).upper())
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        _set_run_font(r, bold=True)
    _mark_content(state)


def _para_heading(doc: Document, text: str, level: int, state: _DocState) -> None:
    clean = _clean_md_text(text)
    tu = clean.upper()
    if tu in STRUCTURAL_HEADS or tu.startswith("ПРИЛОЖЕНИЕ"):
        _request_page_break(state)
        _para_structural(doc, clean, state)
        return
    if re.match(r"^\d+\s", clean) or any(tu.startswith(p) for p in PAGE_BREAK_PREFIXES):
        _request_page_break(state)
    _flush_page_break(doc, state)
    style = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(clean, style=style)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _mark_content(state)


def _para_body(doc: Document, text: str, state: _DocState, indent: bool = True) -> None:
    clean = _clean_md_text(text)
    if not clean:
        return
    _flush_page_break(doc, state)
    p = doc.add_paragraph(clean)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    for r in p.runs:
        _set_run_font(r)
    _mark_content(state)


def _para_caption_table(doc: Document, text: str, state: _DocState) -> None:
    _flush_page_break(doc, state)
    p = doc.add_paragraph(_clean_md_text(text))
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        _set_run_font(r, bold=True)
    _mark_content(state)


def _para_caption_figure(doc: Document, text: str, state: _DocState) -> None:
    _flush_page_break(doc, state)
    p = doc.add_paragraph(_clean_md_text(text))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        _set_run_font(r)
    _mark_content(state)


def _para_code(doc: Document, text: str, state: _DocState) -> None:
    _flush_page_break(doc, state)
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        _set_run_font(r, name="Courier New", size_pt=12)
    _mark_content(state)


def _para_list_item(
    doc: Document,
    text: str,
    state: _DocState,
    *,
    numbered: bool = False,
    number: str | None = None,
) -> None:
    clean = _clean_md_text(text)
    if not clean:
        return
    _flush_page_break(doc, state)
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)
    if numbered and number:
        run = p.add_run(f"{number}. {clean}")
    else:
        run = p.add_run("– " + clean)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_run_font(run)
    _mark_content(state)


def _add_table(doc: Document, rows: list[list[str]], state: _DocState) -> None:
    if not rows:
        return
    _flush_page_break(doc, state)
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell_text = _clean_md_text(row[j] if j < len(row) else "—")
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for par in cell.paragraphs:
                par.paragraph_format.first_line_indent = Cm(0)
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                par.paragraph_format.line_spacing = 1.0
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER if _is_numeric_cell(cell_text) else WD_ALIGN_PARAGRAPH.LEFT
                for r in par.runs:
                    _set_run_font(r, size_pt=12, bold=(i == 0))
    _mark_content(state)


def _is_numeric_cell(text: str) -> bool:
    t = text.strip().replace(",", ".").replace("%", "")
    return bool(re.match(r"^[\d.\-]+$", t))


def _parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def write_docx(md_text: str, path: Path) -> None:
    doc = Document()
    _setup_document(doc)
    _add_page_number_footer(doc)
    state = _DocState()

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[str] = []
    skip_content_outline = False

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()

        if raw.startswith("```"):
            if in_code:
                label = code_buf[0] if code_buf and code_buf[0].startswith("Листинг") else None
                body = code_buf[1:] if label else code_buf
                if label:
                    _para_caption_table(doc, label, state)
                for cl in body:
                    _para_code(doc, cl, state)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        if raw.strip().startswith("|") and "|" in raw[1:]:
            table_buf.append(raw)
            i += 1
            continue
        elif table_buf:
            _add_table(doc, _parse_table_rows(table_buf), state)
            table_buf = []

        if raw.strip() == "---":
            i += 1
            continue

        if raw.startswith("# ") and "ПОЯСНИТЕЛЬНАЯ" in raw.upper():
            p = doc.add_paragraph(_clean_md_text(raw[2:]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                _set_run_font(r, bold=True)
            _mark_content(state)
            i += 1
            if i < len(lines) and lines[i].startswith("к дипломному"):
                p2 = doc.add_paragraph(_clean_md_text(lines[i]))
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p2.paragraph_format.first_line_indent = Cm(0)
                for r in p2.runs:
                    _set_run_font(r)
                _mark_content(state)
                i += 1
            continue

        if raw.startswith("## "):
            t = raw[3:].strip()
            tu = t.upper()
            if tu == "СОДЕРЖАНИЕ":
                _request_page_break(state)
                _para_structural(doc, "СОДЕРЖАНИЕ", state)
                _add_toc(doc, state)
                skip_content_outline = True
                i += 1
                continue
            if skip_content_outline and not raw.startswith("## ВВЕДЕНИЕ"):
                if tu.startswith("ВВЕДЕНИЕ") or re.match(r"^\d", t) or tu == "ЗАКЛЮЧЕНИЕ":
                    skip_content_outline = False
                else:
                    i += 1
                    continue
            _para_heading(doc, t, 1, state)
            i += 1
            continue

        if raw.startswith("### "):
            _para_heading(doc, raw[4:].strip(), 2, state)
            i += 1
            continue

        if raw.startswith("#### "):
            _para_heading(doc, raw[5:].strip(), 3, state)
            i += 1
            continue

        if re.match(r"^\*\*Таблица\s+\d+", raw, re.I) or re.match(r"^Таблица\s+\d+", raw, re.I):
            _para_caption_table(doc, raw, state)
            i += 1
            continue

        if re.match(r"^\*\*Рисунок\s+\d+", raw, re.I) or re.match(r"^Рисунок\s+\d+", raw, re.I):
            i += 1
            continue

        if re.match(r"^Листинг\s+", raw, re.I):
            _para_caption_table(doc, raw, state)
            i += 1
            continue

        num_list = re.match(r"^(\d+)\.\s+(.+)$", raw.strip())
        if num_list:
            _para_list_item(doc, num_list.group(2), state, numbered=True, number=num_list.group(1))
            i += 1
            continue

        if raw.strip().startswith("- "):
            _para_list_item(doc, raw.strip()[2:].strip(), state)
            i += 1
            continue

        if raw.strip():
            _para_body(doc, raw.strip(), state)
        i += 1

    if table_buf:
        _add_table(doc, _parse_table_rows(table_buf), state)

    doc.save(str(path))


def main() -> None:
    md = build_markdown()
    MD_OUT.write_text(md, encoding="utf-8")
    print(f"Written {MD_OUT} ({len(md)} chars)")
    write_docx(md, DOCX_OUT)
    print(f"Written {DOCX_OUT}")


if __name__ == "__main__":
    main()
